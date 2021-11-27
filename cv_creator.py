from docx import Document
from docx.shared import Inches


class CV_Creator:
    def __init__(self, name, mobile_number, email):
        self.name = name
        self.mobile_number = mobile_number
        self.email = email
        self.document = Document()

    def create_header(self):
        self.document.add_picture(
            'Feroz Passport Photo 2.jpg', width=Inches(2.0))
        self.document.add_paragraph(
            self.name + ' | ' + self.mobile_number + ' | ' + self.email
        )

    def create_single_experience(self, experience_para, company, from_date, end_date, job_experience):
        experience_para.add_run(company + '\n').bold = True
        experience_para.add_run(
            from_date + '-' + end_date + '\n').italic = True
        experience_para.add_run(job_experience)

    def create_experiences(self):
        print('------ Starting Experience Section of Resume ------')
        self.document.add_heading('Experiences')
        has_new_Experience = True
        while has_new_Experience:
            experience_para = self.document.add_paragraph()
            company = input('What is the company name? ')
            from_date = input('Start date? ')
            end_date = input('End date? ')
            job_experience = input(
                'What was your experience like in ' + company + ' ')
            self.create_single_experience(
                experience_para, company, from_date, end_date, job_experience)
            user_more_exp_input = input(
                'Would you like to add more experiences? Yes or No ')
            if user_more_exp_input.lower() != 'yes':
                has_new_Experience = False

    def create_single_education_entry(self, education_para, institution, from_date, end_date, institution_experience):
        education_para.add_run(institution + '\n').bold = True
        education_para.add_run(
            from_date + '-' + end_date + '\n').italic = True
        education_para.add_run(institution_experience)

    def create_education_entries(self):
        print('------ Starting Education Section of Resume ------')
        self.document.add_heading('Education')
        has_new_education_entry = True
        while has_new_education_entry:
            education_para = self.document.add_paragraph()
            institute = input('What is the institution name? ')
            from_date = input('Start date? ')
            end_date = input('End date? ')
            institution_experience = input(
                'What was your experience like in ' + institute + '? ')
            self.create_single_education_entry(
                education_para, institute, from_date, end_date, institution_experience)
            user_more_exp_input = input(
                'Would you like to add more education experiences? Yes or No ')
            if user_more_exp_input.lower() != 'yes':
                has_new_education_entry = False

    def save_document(self):
        self.document.save('my_cv.docx')
